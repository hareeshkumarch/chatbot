import io

from docx import Document as DocxDocument


def parse_docx(raw_bytes: bytes) -> list[tuple[str, int | None]]:
    document = DocxDocument(io.BytesIO(raw_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells))
    full_text = "\n".join(paragraphs)
    return [(full_text, None)]
