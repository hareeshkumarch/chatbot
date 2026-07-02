import io

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from app.reports.charts import render_chart_png
from app.reports.models import Report, citation_label

ROUTE_RGB = RGBColor(0x2F, 0x3E, 0xE0)
INK_RGB = RGBColor(0x14, 0x18, 0x1A)
MUTED_RGB = RGBColor(0x5B, 0x64, 0x61)


def render_docx(report: Report) -> bytes:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK_RGB

    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(report.title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = INK_RGB

    if report.subtitle:
        subtitle_paragraph = doc.add_paragraph()
        subtitle_run = subtitle_paragraph.add_run(report.subtitle)
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.color.rgb = MUTED_RGB

    doc.add_paragraph()

    for section in report.sections:
        heading = doc.add_heading(level=1)
        heading_run = heading.add_run(section.heading)
        heading_run.font.color.rgb = ROUTE_RGB
        heading_run.font.size = Pt(16)

        for paragraph_text in section.paragraphs:
            doc.add_paragraph(paragraph_text)

        if section.table and section.table.headers:
            table = doc.add_table(rows=1, cols=len(section.table.headers))
            table.style = "Light Grid Accent 1"
            header_cells = table.rows[0].cells
            for i, header in enumerate(section.table.headers):
                header_cells[i].text = header
            for row_values in section.table.rows:
                row_cells = table.add_row().cells
                for i, value in enumerate(row_values):
                    if i < len(row_cells):
                        row_cells[i].text = value
            doc.add_paragraph()

        if section.chart and section.chart.labels:
            chart_bytes = render_chart_png(section.chart)
            doc.add_picture(io.BytesIO(chart_bytes), width=Inches(6))
            doc.add_paragraph()

    if report.citations:
        source_heading = doc.add_heading(level=1)
        source_heading_run = source_heading.add_run("Sources")
        source_heading_run.font.color.rgb = ROUTE_RGB
        source_heading_run.font.size = Pt(16)
        for citation in report.citations:
            label = citation_label(citation)
            line = doc.add_paragraph()
            run = line.add_run(f"[{citation.get('index')}] {label} \u2014 {citation.get('source_uri', '')}")
            run.font.size = Pt(9.5)
            run.font.color.rgb = MUTED_RGB

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
